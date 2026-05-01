"""
Rollover Module – Audit Automation System

Creates a new financial-year working paper folder by:
  1. Locating the latest previous-year folder inside MAIN ARCHIVE
  2. Copying AUDIT PROGRAMME, AWP, and REPORT into a new CLIENTS folder
  3. Renaming files/folders to the new financial year
  4. Processing Excel files  (year rollover, column shifts, sign-off clearing)
  5. Processing Word documents (year text replacement)
  6. Removing files that are not Excel, Word or PowerPoint
"""

import os
import re
import shutil
import sys
import tempfile
from pathlib import Path

import dropbox
import openpyxl
from docx import Document

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

FOLDERS_TO_COPY = ["AUDIT PROGRAMME", "AWP", "REPORT"]
ALLOWED_EXTENSIONS = {".xlsx", ".xls", ".docx", ".doc", ".pptx", ".ppt", ".xlsm"}

YEAR_PATTERN = re.compile(r"(?:FY\s*)?(\d{4})\s*[-–]\s*(\d{4})", re.IGNORECASE)
SINGLE_YEAR_PATTERN = re.compile(r"(?<!\d)(\d{4})(?!\d)")

# ---------------------------------------------------------------------------
# Dropbox helpers
# ---------------------------------------------------------------------------


def get_dbx_client(access_token: str) -> dropbox.Dropbox:
    return dropbox.Dropbox(access_token)


def list_folders(dbx: dropbox.Dropbox, path: str) -> list[str]:
    """Return immediate sub-folder names under *path*."""
    try:
        result = dbx.files_list_folder(path)
    except dropbox.exceptions.ApiError as exc:
        text = str(exc)
        if "ListFolderError('path'" in text and "not_found" in text:
            raise FileNotFoundError(f"Dropbox path not found: {path}") from exc
        raise
    except dropbox.exceptions.BadInputError as exc:
        text = str(exc)
        if "required scope" in text and "files.metadata.read" in text:
            raise RuntimeError(
                "Dropbox app/token is missing required scope 'files.metadata.read'. "
                "Enable it in Dropbox App Console > Permissions and regenerate the access token."
            ) from exc
        raise
    return [
        e.name
        for e in result.entries
        if isinstance(e, dropbox.files.FolderMetadata)
    ]


def find_latest_year_folder(dbx: dropbox.Dropbox, archive_path: str) -> str | None:
    """Return the latest year-folder name inside archive path."""
    folders = list_folders(dbx, archive_path)
    year_folders = []
    for name in folders:
        try:
            start, end = parse_year_string(name)
        except ValueError:
            continue
        year_folders.append((end, start, name))
    if not year_folders:
        return None
    year_folders.sort(key=lambda t: (t[0], t[1]))
    return year_folders[-1][2]


def resolve_client_archive_path(
    dbx: dropbox.Dropbox,
    archive_base: str,
    client_name: str,
) -> tuple[str, str]:
    """
    Resolve the archive path for a client and return:
      (resolved_archive_path, latest_year_folder)

    Tries common Dropbox structures:
      - <archive_base>/<client_name>
      - <archive_base>/<client_name>/<client_name>
      - <archive_base>
    """
    candidates = []
    for p in (
        f"{archive_base}/{client_name}",
        f"{archive_base}/{client_name}/{client_name}",
        archive_base,
    ):
        clean = re.sub(r"/{2,}", "/", p).rstrip("/")
        if clean and clean not in candidates:
            candidates.append(clean)

    missing_paths: list[str] = []
    no_year_paths: list[str] = []

    for candidate in candidates:
        try:
            latest = find_latest_year_folder(dbx, candidate)
        except FileNotFoundError:
            missing_paths.append(candidate)
            continue

        if latest:
            return candidate, latest
        no_year_paths.append(candidate)

    if no_year_paths:
        joined = ", ".join(no_year_paths)
        raise RuntimeError(
            f"No year folders found under candidate archive paths: {joined}. "
            "Check client name and archive folder structure."
        )

    joined = ", ".join(missing_paths) if missing_paths else ", ".join(candidates)
    raise RuntimeError(
        f"Archive path not found. Checked: {joined}. "
        "Set --archive / DROPBOX_ARCHIVE_BASE to the correct Dropbox root."
    )


def download_folder(
    dbx: dropbox.Dropbox,
    remote_path: str,
    local_dir: Path,
) -> Path:
    """Recursively download *remote_path* into *local_dir* and return the local path."""
    result = dbx.files_list_folder(remote_path, recursive=True)
    for entry in result.entries:
        rel = entry.path_display[len(remote_path):]
        local_path = local_dir / rel.lstrip("/")
        if isinstance(entry, dropbox.files.FolderMetadata):
            local_path.mkdir(parents=True, exist_ok=True)
        elif isinstance(entry, dropbox.files.FileMetadata):
            local_path.parent.mkdir(parents=True, exist_ok=True)
            dbx.files_download_to_file(str(local_path), entry.path_display)
    return local_dir


def upload_folder(
    dbx: dropbox.Dropbox,
    local_dir: Path,
    remote_path: str,
) -> None:
    """Recursively upload *local_dir* to *remote_path* (overwrite existing)."""
    for file_path in local_dir.rglob("*"):
        if file_path.is_file():
            rel = file_path.relative_to(local_dir)
            dbx_path = f"{remote_path}/{rel.as_posix()}"
            dbx.files_upload(
                file_path.read_bytes(),
                dbx_path,
                mode=dropbox.files.WriteMode.overwrite,
                mute=True,
            )


# ---------------------------------------------------------------------------
# Year utilities
# ---------------------------------------------------------------------------


def parse_year_string(year_str: str) -> tuple[int, int]:
    """Parse year text into (start_year, end_year).

    Supported inputs:
      - FY 2023-2024
      - 2023-2024
      - 2024 (interpreted as 2023-2024)
    """
    m = YEAR_PATTERN.search(year_str)
    if m:
        return int(m.group(1)), int(m.group(2))

    single = SINGLE_YEAR_PATTERN.search(year_str)
    if single:
        end_year = int(single.group(1))
        return end_year - 1, end_year

    raise ValueError(f"Cannot parse year from: {year_str}")


def format_year(start: int, end: int, prefix: str = "FY") -> str:
    return f"{prefix} {start}-{end}"


def shift_years_in_text(text: str, old_start: int, old_end: int, new_start: int, new_end: int) -> str:
    """Replace all occurrences of old year range strings with new ones."""
    old_str = str(old_start)
    new_str = str(new_start)
    text = text.replace(str(old_end), str(new_end))
    text = text.replace(old_str, new_str)
    return text


# ---------------------------------------------------------------------------
# Renaming
# ---------------------------------------------------------------------------


def rename_path(path: Path, old_start: int, old_end: int, new_start: int, new_end: int) -> Path:
    """Rename a single file or folder, replacing year references. Returns new path."""
    new_name = path.name
    new_name = new_name.replace(str(old_start), str(new_start))
    new_name = new_name.replace(str(old_end), str(new_end))
    new_name = new_name.replace(str(old_start)[-2:], str(new_start)[-2:])
    new_name = new_name.replace(str(old_end)[-2:], str(new_end)[-2:])
    if new_name != path.name:
        new_path = path.with_name(new_name)
        path.rename(new_path)
        return new_path
    return path


def rename_tree(root: Path, old_start: int, old_end: int, new_start: int, new_end: int) -> None:
    """Rename all folders and files under *root* depth-first."""
    for dirpath, dirnames, filenames in os.walk(root, topdown=False):
        dp = Path(dirpath)
        for f in filenames:
            rename_path(dp / f, old_start, old_end, new_start, new_end)
        renames = []
        for d in dirnames:
            new = rename_path(dp / d, old_start, old_end, new_start, new_end)
            renames.append(new.name)
        dirnames[:] = renames
    rename_path(root, old_start, old_end, new_start, new_end)


# ---------------------------------------------------------------------------
# Excel processing (openpyxl)
# ---------------------------------------------------------------------------


def process_excel(file_path: Path, old_start: int, old_end: int, new_start: int, new_end: int) -> None:
    """Open an Excel file and:
      - Replace old year values with new year values in all cells
      - Move current-year data columns into previous-year columns
      - Clear sign-offs, comments, and metadata cells
    """
    try:
        wb = openpyxl.load_workbook(str(file_path))
    except Exception:
        return

    for ws in wb.worksheets:
        _shift_year_columns(ws, old_start, old_end, new_start, new_end)
        _clear_signoffs_and_metadata(ws)

    wb.save(str(file_path))


def _shift_year_columns(ws, old_start, old_end, new_start, new_end) -> None:
    """Replace year text in cells and shift current-year data to previous-year columns."""
    current_label = str(new_start)
    previous_label = str(old_start)
    header_row = _find_header_row(ws)
    if header_row is None:
        _replace_years_in_cells(ws, old_start, old_end, new_start, new_end)
        return

    current_col = None
    previous_col = None

    for col in range(1, ws.max_column + 1):
        cell_val = str(ws.cell(row=header_row, column=col).value or "")
        if current_label in cell_val and "previous" not in cell_val.lower() and "prior" not in cell_val.lower():
            current_col = col
        elif previous_label in cell_val or "previous" in cell_val.lower() or "prior" in cell_val.lower():
            previous_col = col

    if current_col and previous_col:
        for row in range(header_row + 1, ws.max_row + 1):
            ws.cell(row=row, column=previous_col).value = ws.cell(row=row, column=current_col).value
            ws.cell(row=row, column=current_col).value = None

    _replace_years_in_cells(ws, old_start, old_end, new_start, new_end)


def _replace_years_in_cells(ws, old_start, old_end, new_start, new_end) -> None:
    """Walk every cell and replace old year strings with new ones."""
    for row in ws.iter_rows(min_row=1, max_row=ws.max_row, max_col=ws.max_column):
        for cell in row:
            if cell.value and isinstance(cell.value, str):
                cell.value = shift_years_in_text(cell.value, old_start, old_end, new_start, new_end)


def _find_header_row(ws) -> int | None:
    """Heuristic: find the first row that looks like a header (contains year-like values)."""
    for row in range(1, min(ws.max_row + 1, 20)):
        for col in range(1, ws.max_column + 1):
            val = str(ws.cell(row=row, column=col).value or "")
            if YEAR_PATTERN.search(val):
                return row
    return None


def _clear_signoffs_and_metadata(ws) -> None:
    """Clear cells that contain sign-off or metadata keywords."""
    keywords = [
        "prepared by", "reviewed by", "date", "sign", "initial",
        "checked", "approved", "comment", "remark", " Reviewed",
    ]
    for row in ws.iter_rows(min_row=1, max_row=ws.max_row, max_col=ws.max_column):
        for cell in row:
            if cell.value and isinstance(cell.value, str):
                lower = cell.value.lower()
                if any(k in lower for k in keywords):
                    cell.value = None


# ---------------------------------------------------------------------------
# Word processing (python-docx)
# ---------------------------------------------------------------------------


def process_word(file_path: Path, old_start: int, old_end: int, new_start: int, new_end: int) -> None:
    """Replace old year references with new year references in a Word document."""
    doc = Document(str(file_path))

    for paragraph in doc.paragraphs:
        paragraph.text = shift_years_in_text(
            paragraph.text, old_start, old_end, new_start, new_end
        )

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.text = shift_years_in_text(
                        paragraph.text, old_start, old_end, new_start, new_end
                    )

    for section in doc.sections:
        for header in section.header.paragraphs:
            header.text = shift_years_in_text(
                header.text, old_start, old_end, new_start, new_end
            )
        for footer in section.footer.paragraphs:
            footer.text = shift_years_in_text(
                footer.text, old_start, old_end, new_start, new_end
            )

    doc.save(str(file_path))


# ---------------------------------------------------------------------------
# File cleanup
# ---------------------------------------------------------------------------


def remove_disallowed_files(root: Path) -> None:
    """Delete files whose extension is not in ALLOWED_EXTENSIONS."""
    for file_path in root.rglob("*"):
        if file_path.is_file() and file_path.suffix.lower() not in ALLOWED_EXTENSIONS:
            file_path.unlink()


# ---------------------------------------------------------------------------
# Core orchestration
# ---------------------------------------------------------------------------


def run_rollover(
    dbx: dropbox.Dropbox,
    client_name: str,
    new_financial_year: str,
    archive_base: str = "/MAIN ARCHIVE",
    clients_base: str = "/CLIENTS",
) -> dict:
    """
    Main entry point.  Downloads, processes, and re-uploads the rolled-over
    working papers for *client_name* into *new_financial_year*.

    Returns a summary dict with keys: status, folders_copied, files_processed.
    """
    new_start, new_end = parse_year_string(new_financial_year)

    # 1. Locate the latest previous-year folder
    client_archive, latest_folder = resolve_client_archive_path(dbx, archive_base, client_name)

    old_start, old_end = parse_year_string(latest_folder)

    # 2. Download the relevant folders to a temp directory
    tmp_dir = Path(tempfile.mkdtemp(prefix="rollover_"))
    src_dir = tmp_dir / "source"
    src_dir.mkdir()

    folders_copied = []
    for folder in FOLDERS_TO_COPY:
        remote = f"{client_archive}/{latest_folder}/{folder}"
        try:
            local = src_dir / folder
            download_folder(dbx, remote, local)
            folders_copied.append(folder)
        except dropbox.exceptions.ApiError:
            continue

    if not folders_copied:
        shutil.rmtree(tmp_dir)
        return {"status": "error", "message": "None of the required folders were found"}

    # 3. Process downloaded files
    files_processed = {"excel": 0, "word": 0, "removed": 0}

    # 3a. Rename first (files and folders)
    rename_tree(src_dir, old_start, old_end, new_start, new_end)

    # 3b. Process Excel files
    for xlsx in src_dir.rglob("*.xls*"):
        process_excel(xlsx, old_start, old_end, new_start, new_end)
        files_processed["excel"] += 1

    # 3c. Process Word files
    for docx in src_dir.rglob("*.doc*"):
        process_word(docx, old_start, old_end, new_start, new_end)
        files_processed["word"] += 1

    # 3d. Remove disallowed files
    count_before = sum(1 for p in src_dir.rglob("*") if p.is_file())
    remove_disallowed_files(src_dir)
    count_after = sum(1 for p in src_dir.rglob("*") if p.is_file())
    files_processed["removed"] = count_before - count_after

    # 4. Upload to new CLIENTS location
    new_remote = f"{clients_base}/{client_name}/{format_year(new_start, new_end)}"
    upload_folder(dbx, src_dir, new_remote)

    # 5. Cleanup
    shutil.rmtree(tmp_dir)

    return {
        "status": "ok",
        "source_year": format_year(old_start, old_end),
        "target_year": format_year(new_start, new_end),
        "folders_copied": folders_copied,
        "files_processed": files_processed,
    }


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------


def main():
    import argparse
    import json
    import os

    parser = argparse.ArgumentParser(description="Rollover Module – Audit Automation")
    parser.add_argument("--client", required=True, help="Client name")
    parser.add_argument("--year", required=True, help="New financial year, e.g. 'FY 2025-2026'")
    parser.add_argument("--token", default=os.environ.get("DROPBOX_ACCESS_TOKEN"), help="Dropbox access token (or set DROPBOX_ACCESS_TOKEN env var)")
    parser.add_argument("--archive", default="/MAIN ARCHIVE", help="Path to MAIN ARCHIVE folder")
    parser.add_argument("--clients", default="/CLIENTS", help="Path to CLIENTS folder")
    args = parser.parse_args()

    if not args.token:
        parser.error("Dropbox access token required (--token or DROPBOX_ACCESS_TOKEN env var)")

    try:
        dbx = get_dbx_client(args.token)
        result = run_rollover(dbx, args.client, args.year, args.archive, args.clients)
        print(json.dumps(result))
        if result.get("status") == "error":
            sys.exit(2)
    except Exception as exc:
        print(json.dumps({"status": "error", "message": str(exc)}))
        sys.exit(2)


if __name__ == "__main__":
    main()
